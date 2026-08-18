# -*- coding: utf-8 -*-
"""
DataMeasurement 운영자 매뉴얼 원고(.md) -> Word 문서(.docx) 변환 스크립트.

사용법:
    python _build_operator_manual_docx.py

원고(DataMeasurement_Operator_Manual_v1.0.md)를 수정한 뒤 이 스크립트를 다시 실행하면
DataMeasurement_Operator_Manual_v1.0.docx 가 같은 내용으로 재생성된다.
경로는 이 스크립트 파일 위치를 기준으로 계산하므로 어느 작업 디렉터리에서 실행해도 동작한다.

--------------------------------------------------------------------------------------
원고 문법 (일반 마크다운과 다른 부분만 정리)
--------------------------------------------------------------------------------------
<!--COVER:TITLE-->텍스트         표지 항목 한 줄 (TITLE/SUBTITLE/VERSION/DATE/COMPANY/CONTACT)
<!--PAGEBREAK-->                강제 페이지 나누기
<!--H1-->텍스트                 번호가 없는 제목을 Heading 1 로 삽입 (표지 뒤 "개 요", "목차" 전용)
<!--TOC-->                      Word 자동 목차 필드 삽입 + 안내문 + 이어서 페이지 나누기

N. 제목                         Heading 1  (장, 예: "1. 시작하기 전에")
## N-M. 제목                    Heading 2  (절, 예: "## 1-1. 이 매뉴얼의 대상과 범위")
## N-M-K. 제목                  Heading 3  (항, 예: "## 1-1-1. ...")
부록 X. 제목                    Heading 1  (부록, 예: "부록 A. 용어 설명")

[그림 N-M] 설명                 회색 자리표시자 표(1행1열) + "※ 이 자리에 화면 캡처 이미지를
                                 삽입하세요" 안내문. 정확히 다음 줄의 "> 캡처 대상: ..." 을
                                 표 아래 작은 회색 이탤릭 문단으로 함께 렌더링한다.

- 텍스트                        List Bullet
①②③... 텍스트                  List Number (번호 목록 — 아래 설계 노트 참고)
| a | b |                       Word 표(Table Grid)로 변환 (헤더 행 굵게, 구분선 행은 무시)
**굵게**                        굵은 run

--------------------------------------------------------------------------------------
설계 노트 — 번호 목록에 "1. 2. 3." 대신 "①②③" 을 쓰는 이유
--------------------------------------------------------------------------------------
장(章)/절(節) 제목이 "N. 제목" / "N-M. 제목" 형식이라, 본문에 "1. 클릭합니다" 같은 일반
숫자형 번호 목록을 그대로 쓰면 장 제목 정규식과 줄 시작 패턴이 겹쳐 구분할 수 없다
(예: "1. 로그인 버튼을 클릭합니다." 이 "1. 시작하기 전에" 와 똑같이 "숫자+마침표" 로 시작함).
그래서 원고 본문의 조작 절차는 전부 원 안 숫자(①②③...)로 쓰고, 장/절 제목만 "숫자(-숫자)*."
패턴을 쓰도록 원고 작성 규칙을 분리했다. 이 스크립트는 그 규칙을 그대로 반영해 원 안 숫자로
시작하는 줄을 List Number 스타일로 렌더링하고, "숫자(-숫자)*." 패턴은 항상 제목으로 취급한다.
"""

import io
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----------------------------------------------------------------------------
# 경로 (스크립트 위치 기준 — 어디서 실행해도 동작)
# ----------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_MD = os.path.join(SCRIPT_DIR, 'DataMeasurement_Operator_Manual_v1.0.md')
OUT_DOCX = os.path.join(SCRIPT_DIR, 'DataMeasurement_Operator_Manual_v1.0.docx')

FONT_NAME = '맑은 고딕'
GRAY_FILL = 'D9D9D9'
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)

CIRCLED_NUMBERS = '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳'

COVER_FIELD_STYLE = {
    'TITLE': (30, True),
    'SUBTITLE': (16, False),
    'VERSION': (13, False),
    'DATE': (13, False),
    'COMPANY': (13, False),
    'CONTACT': (13, False),
}

HEADING3_RE = re.compile(r'^(\d+-\d+-\d+)\.\s+(.+)$')
HEADING2_RE = re.compile(r'^(\d+-\d+)\.\s+(.+)$')
HEADING1_RE = re.compile(r'^(\d+)\.\s+(.+)$')
APPENDIX_RE = re.compile(r'^(부록\s*[A-Z])\.\s*(.+)$')
FIGURE_RE = re.compile(r'^\[그림\s*(\d+-\d+)\]\s*(.*)$')
CAPTURE_RE = re.compile(r'^>\s*캡처\s*대상:\s*(.*)$')
COVER_RE = re.compile(r'^<!--COVER:(\w+)-->(.*)$')
BOLD_SPLIT_RE = re.compile(r'(\*\*.+?\*\*)')


# ----------------------------------------------------------------------------
# 폰트 (맑은 고딕 + eastAsia 명시 — 이걸 빠뜨리면 Word 에서 한글이 엉뚱한 폰트로 나온다)
# ----------------------------------------------------------------------------
def apply_east_asian_font(style, font_name):
    style.font.name = font_name  # ascii/hAnsi 지정 (python-docx 기본 동작)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def setup_fonts(doc):
    style_names = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'List Bullet', 'List Number']
    for name in style_names:
        try:
            apply_east_asian_font(doc.styles[name], FONT_NAME)
        except KeyError:
            pass  # 템플릿에 해당 스타일이 없으면 건너뜀 (필수 스타일이 아니면 문제 없음)


# ----------------------------------------------------------------------------
# 인라인 **굵게** 파싱
# ----------------------------------------------------------------------------
def add_runs_with_bold(paragraph, text):
    parts = BOLD_SPLIT_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ----------------------------------------------------------------------------
# 표지
# ----------------------------------------------------------------------------
def add_cover_line(doc, key, text):
    size, bold = COVER_FIELD_STYLE.get(key, (12, False))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if key == 'TITLE':
        p.paragraph_format.space_before = Pt(120)
        p.paragraph_format.space_after = Pt(18)
    else:
        p.paragraph_format.space_after = Pt(6)


# ----------------------------------------------------------------------------
# Word 자동 목차 필드
# ----------------------------------------------------------------------------
def add_toc_field(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    r_element = run._r

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')

    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = '목차를 업데이트하려면 F9 를 누르세요'

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    r_element.append(fld_begin)
    r_element.append(instr_text)
    r_element.append(fld_separate)
    r_element.append(placeholder_text)
    r_element.append(fld_end)


# ----------------------------------------------------------------------------
# 바닥글 페이지 번호 (실패해도 문서 생성 자체는 계속 진행 — 필수 기능 아님)
# ----------------------------------------------------------------------------
def add_page_number_footer(doc):
    try:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        r_element = run._r

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'PAGE'
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        r_element.append(fld_begin)
        r_element.append(instr_text)
        r_element.append(fld_end)
    except Exception as ex:
        sys.stderr.write('[경고] 바닥글 페이지 번호 삽입 실패(무시하고 계속 진행): {0}\n'.format(ex))


# ----------------------------------------------------------------------------
# 그림 자리표시자 (회색 1x1 표 + 캡처 대상 안내 문단)
# ----------------------------------------------------------------------------
def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_figure_placeholder(doc, label, description):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade_cell(cell, GRAY_FILL)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('[그림 {0}] {1}'.format(label, description))
    run1.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('※ 이 자리에 화면 캡처 이미지를 삽입하세요')
    run2.italic = True
    run2.font.color.rgb = GRAY_TEXT

    space_p = doc.add_paragraph()
    space_p.paragraph_format.space_after = Pt(4)


def add_capture_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('캡처 대상: ' + text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_TEXT


# ----------------------------------------------------------------------------
# 마크다운 표 -> Word 표(Table Grid)
# ----------------------------------------------------------------------------
def is_table_separator_row(cells):
    for c in cells:
        if not re.match(r'^:?-+:?$', c.strip()):
            return False
    return True


def split_table_row(line):
    inner = line.strip()
    if inner.startswith('|'):
        inner = inner[1:]
    if inner.endswith('|'):
        inner = inner[:-1]
    return [c.strip() for c in inner.split('|')]


def add_markdown_table(doc, table_lines):
    rows = [split_table_row(ln) for ln in table_lines]
    if len(rows) >= 2 and is_table_separator_row(rows[1]):
        header = rows[0]
        data_rows = rows[2:]
    else:
        header = rows[0]
        data_rows = rows[1:]

    n_cols = len(header)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.paragraphs[0].text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    for row_cells in data_rows:
        row = table.add_row()
        for idx in range(n_cols):
            text = row_cells[idx] if idx < len(row_cells) else ''
            cell = row.cells[idx]
            cell.paragraphs[0].text = ''
            add_runs_with_bold(cell.paragraphs[0], text)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# ----------------------------------------------------------------------------
# 본문 변환 메인 루프
# ----------------------------------------------------------------------------
def strip_markdown_heading_prefix(line):
    return re.sub(r'^#{1,6}\s+', '', line)


def convert(doc, lines):
    figure_count = 0
    i = 0
    n = len(lines)

    while i < n:
        raw_line = lines[i].rstrip('\n')
        stripped = raw_line.strip()

        if stripped == '':
            i += 1
            continue

        # 표지 항목
        cover_m = COVER_RE.match(stripped)
        if cover_m:
            add_cover_line(doc, cover_m.group(1), cover_m.group(2))
            i += 1
            continue

        # 강제 페이지 나누기
        if stripped == '<!--PAGEBREAK-->':
            doc.add_page_break()
            i += 1
            continue

        # 번호 없는 Heading 1 ("개 요", "목차")
        if stripped.startswith('<!--H1-->'):
            doc.add_heading(stripped[len('<!--H1-->'):], level=1)
            i += 1
            continue

        # 자동 목차 필드 + 이어서 페이지 나누기
        if stripped == '<!--TOC-->':
            add_toc_field(doc)
            doc.add_page_break()
            i += 1
            continue

        # 그림 자리표시자 (+ 바로 다음 줄의 캡처 대상 설명)
        fig_m = FIGURE_RE.match(stripped)
        if fig_m:
            label, description = fig_m.group(1), fig_m.group(2)
            add_figure_placeholder(doc, label, description)
            figure_count += 1
            if i + 1 < n:
                cap_m = CAPTURE_RE.match(lines[i + 1].strip())
                if cap_m:
                    add_capture_note(doc, cap_m.group(1))
                    i += 1
            i += 1
            continue

        # (그림 자리표시자 없이 단독으로 나타나는 캡처 대상 줄은 없어야 하지만, 방어적으로 일반 문단 처리)
        cap_only_m = CAPTURE_RE.match(stripped)
        if cap_only_m:
            add_capture_note(doc, cap_only_m.group(1))
            i += 1
            continue

        # 마크다운 표 블록 (연속된 '|' 시작 줄 전부 수집)
        if stripped.startswith('|'):
            table_lines = []
            j = i
            while j < n and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].strip())
                j += 1
            add_markdown_table(doc, table_lines)
            i = j
            continue

        heading_text = strip_markdown_heading_prefix(stripped)

        h3 = HEADING3_RE.match(heading_text)
        if h3:
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        h2 = HEADING2_RE.match(heading_text)
        if h2:
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        appendix_m = APPENDIX_RE.match(heading_text)
        if appendix_m:
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        h1 = HEADING1_RE.match(heading_text)
        if h1:
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        # 불릿 목록
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs_with_bold(p, stripped[2:].strip())
            i += 1
            continue

        # 번호 목록 (원 안 숫자 ①②③... — 위 "설계 노트" 참고)
        if stripped[0] in CIRCLED_NUMBERS:
            p = doc.add_paragraph(style='List Number')
            add_runs_with_bold(p, stripped)
            i += 1
            continue

        # 그 외 일반 문단
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    return figure_count


def main():
    if not os.path.isfile(SRC_MD):
        sys.stderr.write('원고 파일을 찾을 수 없습니다: {0}\n'.format(SRC_MD))
        sys.exit(1)

    with io.open(SRC_MD, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_fonts(doc)
    add_page_number_footer(doc)

    figure_count = convert(doc, lines)

    doc.save(OUT_DOCX)

    paragraph_count = len(doc.paragraphs)
    print('생성 완료: {0} (문단 {1}개, 그림 자리표시자 {2}개)'.format(
        OUT_DOCX, paragraph_count, figure_count))


if __name__ == '__main__':
    main()
